#!/usr/bin/env python3
"""
Generate a reference .docx template for pandoc with styling that matches
the release notes example document.

Style specs (from examples/release-notes-example.docx):
- Title: 28pt, default color (dark/black)
- Subtitle: 14pt, #595959 gray
- Heading 1: 20pt, #0F4761 (teal/dark blue), space_before=18pt, space_after=4pt
- Heading 2: 16pt, #0F4761
- Heading 3: 14pt, #0F4761
- Table headers: fill #44546A, white text (#FFFFFF), 11pt
- Body: Default font (Aptos/Calibri depending on Word version)

Usage:
    pip install python-docx
    python3 generate-reference-docx.py

This creates release-reference.docx in the same directory.

NOTE: The preferred approach is to use the actual example docx
(examples/release-notes-example.docx) as the reference template directly.
This script is a fallback for regenerating the template if needed.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Emu, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def set_style_font(style, size=None, color=None, bold=None):
    """Configure font properties on a style."""
    font = style.font
    if size:
        font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor(*color)
    if bold is not None:
        font.bold = bold


def main():
    doc = Document()

    # -- Title: 28pt (355600 EMU = 28pt) --
    title_style = doc.styles["Title"]
    title_style.font.size = Emu(355600)
    title_style.paragraph_format.space_after = Emu(50800)

    # -- Subtitle: 14pt (177800 EMU), #595959 --
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.size = Emu(177800)
    subtitle_style.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # -- Heading 1: 20pt (254000 EMU), #0F4761 --
    h1 = doc.styles["Heading 1"]
    h1.font.size = Emu(254000)
    h1.font.color.rgb = RGBColor(0x0F, 0x47, 0x61)
    h1.paragraph_format.space_before = Emu(228600)
    h1.paragraph_format.space_after = Emu(50800)

    # -- Heading 2: 16pt (203200 EMU), #0F4761 --
    h2 = doc.styles["Heading 2"]
    h2.font.size = Emu(203200)
    h2.font.color.rgb = RGBColor(0x0F, 0x47, 0x61)
    h2.paragraph_format.space_before = Emu(101600)
    h2.paragraph_format.space_after = Emu(50800)

    # -- Heading 3: 14pt (177800 EMU), #0F4761 --
    h3 = doc.styles["Heading 3"]
    h3.font.size = Emu(177800)
    h3.font.color.rgb = RGBColor(0x0F, 0x47, 0x61)
    h3.paragraph_format.space_before = Emu(101600)
    h3.paragraph_format.space_after = Emu(50800)

    # Add sample content so pandoc can pick up the styles
    doc.add_heading("Release Notes, Project Version X.Y.Z", level=0)  # Title

    subtitle_para = doc.add_paragraph("ENGINEERING RELEASE NOTES (X.Y.Z – X.Y.Z)")
    subtitle_para.style = doc.styles["Subtitle"]

    doc.add_heading("Component Versions", level=1)

    # Sample table with styled header
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Component", "Version"]):
        cell = hdr.cells[i]
        cell.text = text
        # Style header cells: fill #44546A, white text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "44546A")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Emu(139700)  # 11pt

    row = table.rows[1]
    row.cells[0].text = "Example"
    row.cells[1].text = "1.0.0"

    doc.add_paragraph("Main goal of this release.")

    doc.add_heading("Change Summary", level=1)
    doc.add_heading("Feature Area", level=2)
    doc.add_paragraph("Change description.", style="List Bullet")

    doc.add_heading("Issue Summary", level=1)
    doc.add_paragraph("TICKET-123 — Description", style="List Bullet")

    doc.add_heading("Known Issues / Workarounds", level=1)
    doc.add_paragraph("Issue description.", style="List Bullet")

    doc.add_heading("Screenshots & Walkthrough", level=1)

    output = Path(__file__).parent / "release-reference.docx"
    doc.save(str(output))
    print(f"Generated: {output}")


if __name__ == "__main__":
    main()
